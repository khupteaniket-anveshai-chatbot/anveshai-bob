"""
Utility script to read and extract content from DOCX files.
Supports reading full content or sampling first N lines.
"""

import os
from docx import Document


def read_docx_lines(file_path, max_lines=None):
    """
    Read content from a DOCX file and return as lines.
    
    Args:
        file_path (str): Path to the DOCX file
        max_lines (int, optional): Maximum number of lines to read. None for all lines.
    
    Returns:
        list: List of text lines from the document
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(file_path)
    lines = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Only add non-empty lines
            lines.append(text)
            if max_lines and len(lines) >= max_lines:
                break
    
    return lines


def read_docx_sample(file_path, sample_lines=100):
    """
    Read a sample of the first N lines from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        sample_lines (int): Number of lines to sample (default: 100)
    
    Returns:
        dict: Dictionary with 'lines', 'total_sampled', and 'is_complete' keys
    """
    lines = read_docx_lines(file_path, max_lines=sample_lines)
    
    # Check if we read the complete document
    doc = Document(file_path)
    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    is_complete = len(lines) >= total_paragraphs
    
    return {
        'lines': lines,
        'total_sampled': len(lines),
        'is_complete': is_complete,
        'total_paragraphs': total_paragraphs
    }


def read_docx_full(file_path):
    """
    Read complete content from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        list: List of all text lines from the document
    """
    return read_docx_lines(file_path, max_lines=None)


def extract_sections(file_path, max_lines=None):
    """
    Extract content organized by sections (headings).
    Assumes headings are in bold or have specific styles.
    
    Args:
        file_path (str): Path to the DOCX file
        max_lines (int, optional): Maximum number of lines to process
    
    Returns:
        dict: Dictionary with section names as keys and content as values
    """
    doc = Document(file_path)
    sections = {}
    current_section = "Introduction"
    current_content = []
    lines_processed = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        lines_processed += 1
        
        # Check if this is a heading (bold text or heading style)
        is_heading = False
        if paragraph.style.name.startswith('Heading'):
            is_heading = True
        elif paragraph.runs:
            # Check if all runs are bold
            is_heading = all(run.bold for run in paragraph.runs if run.text.strip())
        
        if is_heading:
            # Save previous section
            if current_content:
                sections[current_section] = current_content
            # Start new section
            current_section = text
            current_content = []
        else:
            current_content.append(text)
        
        if max_lines and lines_processed >= max_lines:
            break
    
    # Save last section
    if current_content:
        sections[current_section] = current_content
    
    return sections


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_reader.py <path_to_docx_file> [sample_lines]")
        sys.exit(1)
    
    file_path = sys.argv[1]
    sample_lines = int(sys.argv[2]) if len(sys.argv) > 2 else 100
    
    print(f"Reading sample of {sample_lines} lines from: {file_path}\n")
    
    result = read_docx_sample(file_path, sample_lines)
    
    print(f"Total lines sampled: {result['total_sampled']}")
    print(f"Total paragraphs in document: {result['total_paragraphs']}")
    print(f"Complete document: {result['is_complete']}\n")
    print("=" * 80)
    print("CONTENT SAMPLE:")
    print("=" * 80)
    
    for i, line in enumerate(result['lines'], 1):
        print(f"{i:3d}. {line}")
    
    if not result['is_complete']:
        print("\n" + "=" * 80)
        print(f"... {result['total_paragraphs'] - result['total_sampled']} more lines in document")

# Made with Bob
